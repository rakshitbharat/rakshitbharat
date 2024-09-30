import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

def load_json_data(file_path):
    with open(file_path, 'r') as file:
        return json.load(file)

def add_heading(document, text, level=1, color=RGBColor(52, 152, 219)):
    heading = document.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.color.rgb = color

def add_paragraph(document, text, bold=False, italic=False, color=RGBColor(0, 0, 0)):
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color

def add_list_item(document, text):
    p = document.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

def create_resume(data):
    document = Document()
    
    # Set up styles
    styles = document.styles
    style = styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # Personal Info
    add_heading(document, data['personalInfo']['name'], level=1)
    add_paragraph(document, data['personalInfo']['role'], bold=True)
    add_paragraph(document, f"Email: {data['personalInfo']['email']}")
    add_paragraph(document, f"GitHub: {data['personalInfo']['github']}")
    
    # Professional Summary
    add_heading(document, "Professional Summary", level=2)
    years_of_experience = (datetime.now() - datetime.strptime(data['careerStartDate'], "%Y-%m-%d")).days // 365
    summary = data['professionalSummary'].format(yearsOfExperience=years_of_experience)
    add_paragraph(document, summary)
    
    # Skills
    add_heading(document, "Skills", level=2)
    for category, skills in data['skills'].items():
        add_paragraph(document, f"{category}:", bold=True)
        if isinstance(skills, list):
            add_paragraph(document, ", ".join(skills))
        else:
            add_paragraph(document, skills)
    
    # Professional Experience
    add_heading(document, "Professional Experience", level=2)
    for job in data['professionalExperience']:
        add_paragraph(document, f"{job['role']} at {job['company']}", bold=True)
        add_paragraph(document, job['duration'], italic=True)
        for responsibility in job['responsibilities']:
            add_list_item(document, responsibility)
        document.add_paragraph()  # Add space between jobs
    
    # Additional Experience
    add_heading(document, "Additional Experience", level=2)
    add_paragraph(document, data['additionalExperience']['role'], bold=True)
    for detail in data['additionalExperience']['details']:
        add_list_item(document, detail)
    
    # Education
    add_heading(document, "Education", level=2)
    for edu in data['education']:
        add_paragraph(document, f"{edu['degree']} - {edu['institution']}", bold=True)
        add_paragraph(document, f"{edu['location']}, {edu['duration']}")
    
    # Key Projects
    add_heading(document, "Key Projects", level=2)
    for project in data['keyProjects']:
        add_paragraph(document, project['name'], bold=True)
        add_paragraph(document, f"Role: {project['role']}")
        add_paragraph(document, f"Duration: {project['duration']}")
        add_paragraph(document, f"Technologies: {', '.join(project['technologies'])}")
        add_paragraph(document, "Project Details:", bold=True)
        for detail in project['details']:
            add_list_item(document, detail)
        if 'responsibilities' in project:
            add_paragraph(document, "Responsibilities:", bold=True)
            for resp in project['responsibilities']:
                add_list_item(document, resp)
        if 'achievements' in project:
            add_paragraph(document, "Achievements:", bold=True)
            for achievement in project['achievements']:
                add_list_item(document, achievement)
        document.add_paragraph()  # Add space between projects
    
    # Certificates and Awards
    add_heading(document, "Certificates & Awards", level=2)
    for cert in data['certificates']:
        add_paragraph(document, f"{cert['name']} - {cert['issuer']}")
    for award in data['awards']:
        add_paragraph(document, f"{award['name']} - {award['issuer']}")
    
    # Languages
    add_heading(document, "Languages", level=2)
    for lang in data['languages']:
        add_paragraph(document, f"{lang['language']}: {lang['proficiency']}")
    
    # Interests
    add_heading(document, "Interests", level=2)
    interests_list = []
    for interest in data['interests']:
        if 'technology' in interest:
            interests_list.append(interest['technology'])
        elif 'game' in interest:
            interests_list.append(interest['game'])
        elif 'category' in interest:
            interests_list.extend([activity['name'] for activity in interest['activities']])
        elif 'activity' in interest:
            interests_list.append(interest['activity'])
    add_paragraph(document, ", ".join(interests_list))
    
    # Thank You
    add_heading(document, "Thank You", level=2)
    add_paragraph(document, "Thank you for taking the time to review my resume. I am excited about the opportunity to contribute my skills and experience to your team. I look forward to discussing how I can add value to your organization.")
    add_paragraph(document, "For any further information or to schedule an interview, please don't hesitate to contact me using the information provided at the top of this resume.")
    
    return document

if __name__ == "__main__":
    data = load_json_data('data.json')
    resume = create_resume(data)
    resume.save('Rakshit_Patel_Resume.docx')
    print("Resume created successfully as 'Rakshit_Patel_Resume.docx'")